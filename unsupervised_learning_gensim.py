# -*- coding: utf-8 -*-
"""
Created on Sun May  3 00:02:29 2020

@author: cvicentm
"""

# -*- coding: utf-8 -*-
"""
Created on Fri Apr 17 07:46:39 2020

@author: cvicentm
"""
from sklearn.decomposition import LatentDirichletAllocation
#import pyLDAvis.sklearn
import pickle
from create_corpus import create_corpus
import timeit
import datetime

import gensim, spacy, logging, warnings
import gensim.corpora as corpora
from gensim.utils import lemmatize, simple_preprocess
from gensim.models import CoherenceModel
from PIL import ImageColor
from create_corpus import normalize_word
from docx import Document
from docx.shared import RGBColor
from nltk import word_tokenize
from random import randint
import os
from tqdm import tqdm
import matplotlib.pyplot as plt
import numpy as np
from datetime import datetime

name_log_file = datetime.now().strftime('logs\load_subtitles_%d_%m_%Y.log')
    
logging.basicConfig(filename=name_log_file, level=logging.WARNING, 
                    format="%(asctime)s:%(filename)s:%(lineno)d:%(levelname)s:%(message)s")



def printColorWordDocument(number,colors,generator_normalize,dic_subtitles,lda_model,corpus):
    #make an explanation of every of the parameters of this function
    corp_cur = corpus[number] 
    topic_percs, wordid_topics, wordid_phivalues = lda_model[corp_cur]
    word_dominanttopic = [(lda_model.id2word[wd], topic[0]) for wd, topic in wordid_topics if topic]
                          
    
    def wordToDocument (word, color_hex):
        color_rgb = ImageColor.getrgb(color_hex)
        run=paragraph.add_run(word+" ")
        font = run.font
        font.color.rgb = RGBColor(color_rgb[0], color_rgb[1], color_rgb[2])
    
    document_classified=[]
    word_dominanttopic_dict=dict(word_dominanttopic)
    dict_one_subtitle_token=word_tokenize(dic_subtitles[list(dic_subtitles.keys())[number]])
    for word in dict_one_subtitle_token:
        topic_word=len(colors)-1
        if normalize_word(word) in generator_normalize[number]:
            try:
                topic_word=word_dominanttopic_dict[normalize_word(word)]
            except KeyError as error:
                # Escribir aquí un log en vez de un print
                logging.warning("OSError --- key error")
                
        else:
            topic_word=len(colors)-1
        document_classified.append((word,topic_word))
    
    
    #document_classified=[(word,word_dominanttopic_dict[word]) for word in generator_normalize[0]]
    document = Document()
    paragraph = document.add_paragraph()
    [wordToDocument(word,colors[topic]) for word,topic in document_classified]
    
    word_subtitles_colors="word"+list(dic_subtitles.keys())[number][9:-1]+".docx"
    document.save(word_subtitles_colors)
    

def training_model(n_documents,n_topics,id2word, corpus):
    #ESCRIBIR QUE HACE ESTA FUNCIÓN Y PARA QUE SIRVE CADA UNO DE SUS PARÁMETROS
    
    
    print("the model is being trained with: "+str(n_topics)+ "topics")
    
    lda_model = gensim.models.ldamodel.LdaModel(corpus=corpus,
                                               id2word=id2word,
                                               num_topics=n_topics, 
                                               random_state=100,
                                               update_every=1,
                                               chunksize=100,
                                               passes=10,
                                               alpha='auto',
                                               per_word_topics=True)
    
    
    #tengo bastantes dudas si en el parámetro texts tengo que poner generator_normalize o no
    coherencemodel = CoherenceModel(model=lda_model, corpus=corpus, dictionary=id2word, coherence='u_mass')
    coherence_values = coherencemodel.get_coherence()
    #the model is gonna be saved
    # if the number of subtitles doest change, we can use the same model than the last time
    
    file_lda_model = 'pickle\lda_model_'+str(n_topics)+'_'+str(n_documents)+'.sav'
    pickle.dump(lda_model, open(file_lda_model, 'wb'))
    
    return coherence_values

#PROGRAM......................................................................
def LDAmodel( n_topics, n_documents, n_printedDocuments, step=1, start=1):   
    #Tengo que escribir para que sirve cada cosa que hace el gensim
    coherencemodelArray=[]
    try: 
        generator_normalize = pickle.load(open("pickle\generator_normalize_"+str(n_documents)+".txt", "rb"))
        dic_subtitles = pickle.load(open("pickle\dic_subtitles_"+str(n_documents)+".pickle", "rb"))
        #[generator_normalize, dic_subtitles]=create_corpus(n_documents)
        id2word = pickle.load(open("pickle\id2word_"+str(n_documents)+".txt", "rb"))
        corpus = pickle.load(open("pickle\corpus_"+str(n_documents)+".txt", "rb"))
        print("generator_normalize, id2word and corpus has been imported")
    except IOError:
        print(" the corpus with: "+str(n_topics)+" is being created...")
        [generator_normalize, dic_subtitles]=create_corpus(n_documents)
        file_generator_normalize = 'pickle\generator_normalize_'+str(n_documents)+'.txt'
        pickle.dump(generator_normalize, open(file_generator_normalize, 'wb'))
        
        file_dic_subtitles = 'pickle\dic_subtitles_'+str(n_documents)+'.pickle'
        pickle.dump(dic_subtitles, open(file_dic_subtitles, 'wb'))
        #this is creating a dictionary with all de different words of the document
        id2word = corpora.Dictionary(generator_normalize)
        file_id2word = 'pickle\id2word_'+str(n_documents)+'.txt'
        pickle.dump(id2word, open(file_id2word, 'wb'))
        # Create Corpus: Term Document Frequency
        corpus = [id2word.doc2bow(text) for text in generator_normalize]
        file_corpus = 'pickle\corpus_'+str(n_documents)+'.txt'
        pickle.dump(corpus, open(file_corpus, 'wb'))
        
        print("Proccess of creating corpus has ended")
    
    for n_topics in range(start, n_topics, step):
        file_lda_model = 'pickle\lda_model_'+str(n_topics)+'_'+str(n_documents)+'.sav'
        try:
           
            f=open(file_lda_model, 'rb')
            lda = pickle.load(f)
            print("The model has been trained previously with..."+str(n_topics)+" n_topics")            
            coherencemodel = CoherenceModel(model=lda, corpus=corpus, dictionary=id2word, coherence='u_mass')
            coherence_values = coherencemodel.get_coherence()
            coherencemodelArray.append(coherence_values)
            
        except IOError:
            
            print("FINALLY: the LDA model has to be trained for "+str(n_documents)+" n_documents and "+str(n_topics)+" n_topics, trained")
            
            tic_all_processing=timeit.default_timer()
            #function based on : https://www.machinelearningplus.com/nlp/topic-modeling-gensim-python/#13viewthetopicsinldamodel
            coherencemodelArray.append(training_model(n_documents,n_topics,id2word,corpus))                
            toc_all_processing=timeit.default_timer()
            try: 
                time_lda_fit=str(datetime.timedelta(seconds=int(float(toc_all_processing-tic_all_processing))))
                print("The process of training lda model with "+str(n_topics)+" n_topics and "+str(n_documents)+" n_documents, has taken "+time_lda_fit+" seconds")    
            except AttributeError: 
                print("The process of training lda model with "+str(n_topics)+" n_topics and "+str(n_documents)+" n_documents, has ended")
                
    
    x = range(start, n_topics+1, step)
    #n_topics+1 because has to have the same weight than coherencemodelArray
    plt.plot(x, coherencemodelArray)
    plt.xlabel("N_Topics")
    plt.ylabel("Coherence")
    plt.legend(("coherence_values"), loc='best')
    plt.show()
    
    best_n_topic=coherencemodelArray.index(min(coherencemodelArray))+start
    print("el mejor modelo es: "+'pickle\lda_model_'+str(best_n_topic)+'_'+str(n_documents)+'.sav')
    f=open('pickle\lda_model_'+str(best_n_topic)+'_'+str(n_documents)+'.sav', 'rb')
    lda = pickle.load(f)
    document_per_topic=list(lda.get_document_topics(corpus))
    """
    corp_cur = corpus[1]
    topic_percs, wordid_topics, wordid_phivalues = lda[corp_cur]
    print(wordid_topics)
    """
    array_topic_per_document = np.zeros((len(document_per_topic), best_n_topic))
    
    
    for i in range(len(document_per_topic)):
        for j in range(len(document_per_topic[i])):
            try:    
                array_topic_per_document[i][document_per_topic[i][j][0]]= document_per_topic[i][j][1]
            except IndexError:
                #EN ESTE LOG sería necesario ponerle, cual ha sido el subtítulo que ha dado problemas e identifcar porque
                logging.warning("array_topic_per_document out of range in position n_document: "+str(i)+" and topic: "+str(j)+" \n")
    #NUMBER OF DOCUMENTs to print results on word
    
    
    return array_topic_per_document, best_n_topic, dic_subtitles,lda,generator_normalize,corpus,id2word,coherencemodelArray
"""
print("se va a proceder a guardar el modelo en un fichero")
# if the number of subtitles doest change, we can use the same model than the last time
pickle.dump(lda_model, open(file_lda_model, 'wb'))
#creation of the matrix of colors to print documents
colors = []
    
for i in range(N_TOPICS):
    colors.append('#%06X' % randint(0, 0xFFFFFF))
colors.append('#000000')
#creation of the directory which content all documents printed
if not os.path.exists('word'):
    os.makedirs('word')  

print("colour´s documented are being printed")
for i in tqdm(range(n_printedDocuments)):
    printColorWordDocument(i,colors,generator_normalize,dic_subtitles,lda_model,corpus)
"""

""""
pip install pyldavis==2.1.0
print("estamos con el tema de imprimir las gráficas")
pyLDAvis.enable_notebook()
panel = pyLDAvis.sklearn.prepare(lda, Bow_matrix, vectorizer, mds='tsne')
pyLDAvis.display(panel)
"""


    
#Sería interesante hacer clustering con la matriz array_topic_per_document, consiguiendo saber que documentos se parecen y cuales no

