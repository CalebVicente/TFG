# -*- coding: utf-8 -*-
"""
Created on Wed May  6 17:26:56 2020

@author: cvicentm
"""
#EN ESTE PROGRAMA ES IMPORTANTISIO CAMBIAR EL NOMBE DE VARIABLES Y REORDENARLO, PORQUE MUCHO ESTÁ COPIADO Y PEGADO
from sklearn.decomposition import LatentDirichletAllocation
#import pyLDAvis.sklearn

from create_corpus import create_corpus

import gensim
import gensim.corpora as corpora
from gensim.utils import lemmatize, simple_preprocess
from gensim.models import CoherenceModel
from PIL import ImageColor
from create_corpus import normalize_word
from docx import Document
from docx.shared import RGBColor
from nltk import word_tokenize
from random import randint
import os
from tqdm import tqdm
import matplotlib.pyplot as plt
import numpy as np


import pandas as pd
from sklearn.cluster import KMeans
import get_data

from unsupervised_learning_gensim import LDAmodel
from unsupervised_learning_gensim import printColorWordDocument

N_TOPICS =80
#este parámetro no se puede añadir a mano
n_printedDocuments =5
max_clusters=100

[files, max_documents] = get_data.get_NameFiles()

#if we want to change the number of documents to analized we can do it here
n_documents=max_documents
def validator_cluster(array_topic_per_document, min_cluster=1, max_cluster=n_documents):
	"""This function is going to take the percentaje of the topic of every document, and will validate what number of
	clusters group best similar documents refers to topics
		-min_cluster: minimum number of cluster to group the documents
		-max_cluster: maximum number of cluster to group the documents (this value cant be higher than than number of documents)"""
	print("validating number of clusters...")
	Number_clusters = range(min_cluster, max_cluster)
	#existen muchisimas variables que se puden cambiar, y que probablemente haya que parametrizar, y probablemente validar
	#darle un buen repaso a este tema
	kmeans = [KMeans(n_clusters=i) for i in tqdm(Number_clusters)]
	kmeans
	score = [kmeans[i].fit(array_topic_per_document).score(array_topic_per_document) for i in range(len(kmeans))]
	
	return score

def topic_per_document_pandas(array_topic_per_document, best_n_topic, dic_subtitles):
    
    columns=[]
    
    for j in range(np.shape(array_topic_per_document)[0]):
        columns.append(list(dic_subtitles.keys())[j])
    
    
    clusters = np.zeros(len(list(dic_subtitles.keys())))
    acc = 0
    for subtitle in list(dic_subtitles.keys()):
        for i in range(len(index_clusters)):
            if subtitle in index_clusters[i]:
                clusters[acc]=i
        acc=acc+1
    
    title=[]
    #title.append("clusters")
    for i in range(best_n_topic):
        title.append('Topic_'+str(i))


    dataframe = pd.DataFrame(array_topic_per_document.T, dtype="str", index=title)
    dataframe.columns=columns
    dataframe = dataframe.T
    dataframe.insert(0,"clusters",clusters)
    dataframe = dataframe.T
    return dataframe

def printClusterDf(dataframe, n_documents, index_clusters):
    """Función que imprime en un excel los documentos pertenecientes a un cluster, y los tópicos a los que pertence"""
    with pd.ExcelWriter('results\\ClusterDf_'+str(n_documents)+'.xlsx') as writer:
        for i in range(len(index_clusters)):
            cluster = index_clusters[i]
            dataframe[cluster].T.astype(float).round(3).to_excel(writer, sheet_name='Cluster'+str(i))

def printDayDf(day, dataframe, n_documents, index_clusters, dic_subtitles):
    """Función que imprime en un excel los documentos pertenecientes a un dia, el cluster al que pertenecen y sus datos"""
    subtitles_day = [subtitle for subtitle in list(dic_subtitles.keys()) if day in subtitle]
    with pd.ExcelWriter('results\\day_'+str(n_documents)+'.xlsx') as writer:
        dataframe[subtitles_day].T.astype(float).round(3).to_excel(writer, sheet_name=day)            

def knee_locator_k_means(score):
	"""This funtion localize where is the optimal number of clusters"""
	from kneed import KneeLocator

	x = range(1, len(score)+1)
	#son super importantes las variables curve y direction o el KneeLocator no funcionará correctaente
	kn = KneeLocator(x, score, curve='concave', direction='increasing')
	
	return kn.knee

def graphic_k_means_validator(knee,score):
	"""This function print a graph score of all the validators scores"""
	import matplotlib.pyplot as plt

	x = range(1, len(score)+1)
	plt.xlabel('number of clusters k')
	plt.ylabel('Sum of squared distances')
	plt.plot(x, score, 'bx-')
	plt.vlines(knee, plt.ylim()[0], plt.ylim()[1], linestyles='dashed')

def showGraphsLDATrainedInTerminal(dic_subtitles, array_topic_per_document, best_n_topic, n_documents=n_documents):
	"""This function start with the first document an finalize with the las document indicated"""
	for i in range(n_documents):
	    fig, ax = plt.subplots()   
	    ax.plot(np.arange(0,best_n_topic,1),array_topic_per_document[i],'-*',label=list(dic_subtitles.keys())[i])
	    ax.legend()

def showOneGraphLDATrainedInTerminal(name_document, array_topic_per_document, best_n_topic):
    """This function print into a doc document only one subtitles selected by a user"""
    find_slash = name_document.find("\\")
    name_document = name_document[:find_slash]+"\\"+name_document[find_slash+1:]
    number_document = list(dic_subtitles.keys()).index(name_document)

    fig, ax = plt.subplots()   
    ax.plot(np.arange(0,best_n_topic,1),array_topic_per_document[number_document],'-*',label=name_document)
    ax.legend()

def similar_subtitles(dic_subtitles,k_means):
    """This function group all the subtitles titles for cluster group"""
    k_means_label = k_means.labels_
    index_clusters=[]
    list_subtitles=list(dic_subtitles.keys())
    k_means_label = list(k_means_optimized.labels_)
    for i in range(knee):
        index = []
        index = [list_subtitles[document_number] for document_number, n_cluster in enumerate(k_means_label) if n_cluster == i]
        index_clusters.append(index)
    
    return index_clusters

def printClusters2Document(index_clusters,n_documents,dic_subtitles):
    """This function put in a document all the subtitles divided on clusters"""
    if not os.path.exists('results'):
        os.makedirs('results')
    file = "results\\clusters_"+str(n_documents)+".txt"
    with open(file, 'w') as f:
        acc = 0
        for cluster in index_clusters:
            f.write("\n")
            f.write("----------------------------------------")
            f.write("N CLUSTER: "+str(acc))
            f.write("----------------------------------------\n")
            
            for subtitles in cluster:
                find_slash = subtitles.find("\\")
                name_document = subtitles[:find_slash]+"\\"+subtitles[find_slash+1:]
                number_document = list(dic_subtitles.keys()).index(name_document)
                f.write("CLUSTER = "+str(acc))
                f.write(" %s\n" % subtitles)
                f.write(str(list(array_topic_per_document[number_document])))
                f.write("\n")
            acc = acc + 1  
        f.close()
def printResults2Document(max_documents, n_documents, dic_subtitles, N_TOPICS, best_n_topics, n_clusters, n_optimized_k_means  ):
    """This function put in a document most important parameters into a document"""
    from datetime import datetime
    
    n_news=len(list(dic_subtitles.keys()))
    
    if not os.path.exists('results'):
        os.makedirs('results')
    
    hour=datetime.now().strftime("%d_%m_%Y.txt")
    file = "results\\results_"+str(n_documents)+"_"+hour
    with open(file, 'w') as f:
        f.write(datetime.now().strftime("%d_%m_%Y"))
        f.write("-------------------------------------------------")
        f.write("\n")
        f.write("Numero de documentos disponibles: ")
        f.write(str(max_documents))
        f.write("\n")
        f.write("Numero de documentos usados: ")
        f.write(str(n_documents))
        f.write("\n")
        f.write("Numero de telediarios: ")
        f.write(str(n_news))
        f.write("\n")
        f.write("Numero de validaciones usadas en el LDA: ")
        f.write(str(N_TOPICS))
        f.write("\n")
        f.write("Número optimo de topicos LDA: ")
        f.write(str(best_n_topics))
        f.write("\n")
        f.write("Número de validaciones de K-means: ")
        f.write(str(n_clusters))
        f.write("\n")
        f.write("Número optimo de clusters en k-means: ")
        f.write(str(n_optimized_k_means))
        f.write("\n")
        f.write("-------------------------------------------------")
        f.write("\n")
        f.write("Tamaño máximo del vocabulario: ")
        f.write(str(len(list(dict(id2word).keys()))))
        f.write("\n")
        f.close()

#PROGRAM-----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

[array_topic_per_document, best_n_topic, dic_subtitles,lda,generator_normalize,corpus,id2word,coherenceModelArray]=LDAmodel(n_topics=N_TOPICS,n_documents=n_documents, n_printedDocuments=n_printedDocuments)

score = validator_cluster(array_topic_per_document, min_cluster=1, max_cluster=max_clusters)

knee = knee_locator_k_means(score)

graphic_k_means_validator(knee,score)

k_means_optimized = KMeans(n_clusters=knee).fit(array_topic_per_document)

showGraphsLDATrainedInTerminal(dic_subtitles, array_topic_per_document, best_n_topic, n_documents=5)

index_clusters = similar_subtitles(dic_subtitles,k_means_optimized)

printClusters2Document(index_clusters,n_documents,dic_subtitles)

#print report about main parameters of the analysis
printResults2Document(max_documents, n_documents, dic_subtitles, N_TOPICS, best_n_topic, max_clusters, knee)

topic_dataframe = topic_per_document_pandas(array_topic_per_document, best_n_topic, dic_subtitles)

printClusterDf(topic_dataframe, n_documents,index_clusters)
#IMPRIMIR DOCUMENTOS DE WORD--------------------------------------------------------------

print("ha llegado a entrenar el modelo")
lda_model = gensim.models.ldamodel.LdaModel(corpus=corpus,
                                               id2word=id2word,
                                               num_topics=best_n_topic, 
                                               random_state=100,
                                               update_every=1,
                                               chunksize=100,
                                               passes=10,
                                               alpha='auto',
                                               per_word_topics=True)


"""
colors = []

print("ha llegado hasta la parte de los colores")   

for i in range(best_n_topic):
    colors.append('#%06X' % randint(0, 0xFFFFFF))
colors.append('#000000')
#creation of the directory which content all documents printed
if not os.path.exists('word\\'+str(n_documents)):
        os.makedirs('word\\'+str(n_documents))
      
print("colour´s documented are being printed")
for i in tqdm(range(n_printedDocuments)):
    printColorWordDocument(i,colors,generator_normalize,dic_subtitles,lda_model,corpus,n_documents)

print("el mejor tópicoooooooooooo:"+str(best_n_topic))
"""

"""with pd.ExcelWriter('output.xlsx') as writer:  
    dataframe.to_excel(writer, sheet_name='Sheet_name_1')"""

"""
import pyLDAvis.gensim
print("estamos haciendo el pyLDAvis")
pyLDAvis.enable_notebook()
data = pyLDAvis.gensim.prepare(lda_model, corpus, id2word)
data
"""